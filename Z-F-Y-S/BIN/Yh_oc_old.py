#!/usr/bin/python3.7
##########    form yahoo open close giornaliero (non continuo) #######################
#######         plot trend in aumento su 3 anni e su 3 mesi            ###############



# Inserire un coefficiente di volatilità
# Inserire trend logaritmici ed esponenziali per valutare la crescita



import pandas_datareader.data as web
import pandas as pd
import numpy as np
from sklearn.linear_model import LinearRegression
import time
from datetime import datetime
import matplotlib.pyplot as plt
#from iex import reference
import datetime as dt
import matplotlib.dates as mdates
import matplotlib.ticker as mticker
plt.style.use('ggplot')
import sys
import os
import shutil
#%matplotlib inline
import warnings
warnings.filterwarnings('ignore')
import mammoth
#pip3.7 install --user python-docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pickle
import h5py
import dateutil.relativedelta
import dateutil.relativedelta as datdelta

#

print('START!!!')
# mi assicuro che la working directory sia quella giusta

try:
    # se sto lavorando in python anywhere non dara errore
    os.chdir("/home/magicpy/Z-F-Y-S")
except:
    pass
print(os.getcwd())

prev_folder = os.path.dirname(os.getcwd())
temp_folder = os.path.join(prev_folder,"mysite","templates", "../todaytrade.html")
print('Save html file in: ', temp_folder)

def yeardelta(date,delta):

    try:
        format_str = '%Y-%m-%d'
        date = datetime.datetime.strptime(date, format_str)
    except:
        pass

    y1 = date.year+delta

    return date.replace(day=date.day, month=date.month, year=y1)

def monthdelta(date, delta):

    try:
        format_str = '%Y-%m-%d'
        date = datetime.datetime.strptime(date, format_str)
    except:
        pass

    m, y = (date.month+delta) % 12, date.year + ((date.month)+delta-1) // 12
    if not m: m = 12
    d = min(date.day, [31,
        29 if y%4==0 and not y%400==0 else 28,31,30,31,30,31,31,30,31,30,31][m-1])
    return date.replace(day=d,month=m, year=y)


#
#

class Yahoo_Scan():

    def __init__(self,stock,in_date,interv,name_ist_str):

        self.stock = stock

        try:

            df_base = web.get_data_yahoo(stock,in_date,interval=interv).copy()

            df = df_base.copy()

            df['datestamp']=pd.to_datetime(df.index)
            df['timestamp']=df.datestamp.apply(lambda x: int(time.mktime(x.timetuple())))
            df["day_perf"]=((df.Close/df.Open)-1)   # today delta open close positivo se in aumento  (ultimo/precedente)-1
            df["day_perf_x100"] = pd.Series(["{0:.2f}%".format(val * 100) for val in df["day_perf"]], index = df.index) #@


            df['Open_p']=df.Open.shift(1)
            df['Close_p']=df.Close.shift(1)
            df['Open_p2']=df.Open.shift(2)
            df['Close_p2']=df.Close.shift(2)
            df['Open_p5']=df.Open.shift(5)
            df['Close_p5']=df.Close.shift(5)
            df['Open_p10']=df.Open.shift(10)
            df['Close_p10']=df.Close.shift(10)

            df['Volume_p']=df.Volume.shift(1)


            df["ema12"] = df["Adj Close"].rolling(window=12).mean()   #E' corretto ma non torna con Etoro
            df["ema26"] = df["Adj Close"].rolling(window=26).mean()   #E' corretto ma non torna con Etoro
            df["macd"] = df.ema12 - df.ema26
            df["signal"] = df["macd"].rolling(window=9).mean()
            df["histogram"] = df.macd - df.signal

            #
            # df['DVolume']=(df.Volume-df.Volume_p)/df.Volume   # Sbagliato bisogna mettere /df.volume_p
            df['DVolume']=(df.Volume/df.Volume_p)-1                         # significa un delta di volume nella notte
            df["ema12v"] = df["DVolume"].rolling(window=12).mean()
            df["ema26v"] = df["DVolume"].rolling(window=26).mean()
            df["macdv"] = df.ema12v - df.ema26v
            df["signalv"] = df["macdv"].rolling(window=9).mean()
            df["histogramv"] = df.macdv - df.signalv
            #
            df['Overnight']=round((df.Open-df.Close_p)/df.Close_p*100+100,2)    ## Overnight apertura-chiusura-ieri      sempre positivo 101 o 99
            df['Dgain']=round((df.Close-df.Open)/df.Open*100+100,0)             ## Dgain     chiusura-aperturamattina    sempre positivo 101 o 99
            df['Tgain_']=df.Dgain.shift(-1)                                     ## Tgain     Dgain di domani
            df['Dgain_p']=df.Dgain.shift(1)                                     ## Dgain_p   Dgain di ieri
            df['Dgain_p2']=df.Dgain.shift(2)                                    ## Dgain_p2  Dgain di ieri l'altro
            df['Dgain_p5']=df.Dgain.shift(5)                                    ## Dgain_p5  Dgain 5gg fà
            df['Dgain_p10']=df.Dgain.shift(10)                                  ## Dgain_p10 Dgain 10gg fa

            df['MDgain_p1']=round((df.Close-df.Open_p)/df.Open_p*100+100,0)
            df['MDgain_p2']=round((df.Close-df.Open_p2)/df.Open_p2*100+100,0)
            df['MDgain_p5']=round((df.Close-df.Open_p5)/df.Open_p5*100+100,0)
            df['MDgain_p10']=round((df.Close-df.Open_p10)/df.Open_p10*100+100,0)


            df['DVolume']=(df.Volume-df.Volume_p)/df.Volume_p*100               ## Dvolume   oggi-ieri
            df["histogram_p"] = df.histogram.shift(1)

            df["day_perf"]=df['Dgain']                                          ## sono comunque calcolate uguali



            # devi rimuovere la prima e l'ultima riga quando fai il rolling
            df['hist_trend'] = 0
            df['hist_peack'] = 0
            df['hist_trend_p'] = 0
            df['hist_changepoint'] = 0

            try:

                ##         /\     23        POSITIVE-ENCREASING 1 POSITIVE-DECREASING 2
                ##        /  \   1  4       NEGATIVE-DECREASING 3 NEGATIVE-DECREASING 3
                ##

                def hist_trend_f(row):
                    '''
                    1 if lower trough
                    0 in between situation
                    -1 if upper peak
                    '''
                    try:
                        if row['histogram'] > 0:
                            if row['histogram'] > row['histogram_p']:
                                val = 2
                            if row['histogram'] <= row['histogram_p']:
                                val = 3
                        elif row['histogram'] < 0:
                            if row['histogram'] > row['histogram_p']:
                                val = 1
                            if row['histogram'] <= row['histogram_p']:
                                val = 4
                        else:
                            val = 0
                        return int(val)
                    except:
                        return int(0)

                def hist_peack_f(row):
                    '''
                    1 if lower trough
                    0 in between situation
                    -1 if upper peak
                    '''
                    try:
                        if (row['hist_trend_p'] == 4) & (row['hist_trend'] == 1):
                            val =1
                        elif (row['hist_trend_p'] == 2) & (row['hist_trend'] == 3):
                            val =-1
                        else:
                            val = 0
                        return val
                    except:
                        print('error hist_peack_f def')
                        return 0


                def hist_changepoint_f(row):
                    '''
                    1 if lower trough
                    0 in between situation
                    -1 if upper peak
                    '''

                    try:
                        if (row['hist_trend_p'] == 1) and (row['hist_trend'] == 2):
                            val =1
                        elif (row['hist_trend_p'] == 3) and (row['hist_trend'] == 4):
                            val =-1
                        else:
                            val = 0
                        return val
                    except:
                        print('error hist_changepoint_f def')
                        return 0


                #########  CREA INDICATORI PER GLI ISTOGRAMMI

                df['hist_trend'] =  df.apply(hist_trend_f, axis=1)
                pd.to_numeric(df['hist_trend'], errors='coerce')

                df['hist_trend_p']= df.hist_trend.shift(1)
                pd.to_numeric(df['hist_trend_p'], errors='coerce')

                df['hist_peack'] =  df.apply(hist_peack_f, axis=1)
                pd.to_numeric(df['hist_peack'], errors='coerce')

                df['hist_changepoint'] = df.apply(hist_changepoint_f, axis=1)
                pd.to_numeric(df['hist_changepoint'], errors='coerce')

                ################### NUOVI INDICATORI CON GIORNI DAL PICCO HISTOGRAMMA ##################################

                df['maxday']=0
                df['daytype']=0
                df['deltaday']=0

                def day_type(row):

                    if row['hist_changepoint'] == 1:
                        val = 1
                    elif row['hist_changepoint'] == -1:
                        val = -1
                    else:
                        val = np.nan
                    return val

                df['daytype'] = df.apply(day_type, axis=1).fillna(method='ffill').fillna(value=0)

                def day_maxday(row):

                    selecteday = row['datestamp']

                    maxday = df[df.datestamp<=selecteday][df.hist_changepoint != 0]['datestamp'].max()

                    return maxday

                df['daytype'] = df.apply(day_type, axis=1).fillna(method='ffill').fillna(value=0)

                df['maxday'] = df.apply(day_maxday, axis=1).astype('datetime64[ns]').fillna(method='bfill')

                df['deltaday'] = (((df.datestamp-df.maxday)/ np.timedelta64(1,'D')).astype(int)+1)*df.daytype



            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                print(self.stock,"Failed at Hisogram indicator: ", e, exc_type, fname, exc_tb.tb_lineno)


            # relevant DGain, Dgain_p, Dgain_p2,

            ##### LINEAR REGRESSION ######################################

            def linear_regress(df_):

                xt = df_.index.tolist()

                x = np.array(df_.timestamp).reshape((-1, 1))

                y_h = np.array(df_.High)
                model = LinearRegression()
                model.fit(x, y_h)
                y_h_pred = model.predict(x)

                y_l = np.array(df_.Low)
                model = LinearRegression()
                model.fit(x, y_l)
                y_l_pred = model.predict(x)

                y_c = np.array(df_.Close)
                model = LinearRegression()
                model.fit(x, y_c)
                y_c_pred = model.predict(x)
                coeff_close = model.coef_[0]           #--->

                return xt, y_h, y_h_pred, y_l, y_l_pred, y_c, y_c_pred, coeff_close

            #####################################

            ##  Linear regression all months
            xt, y_h, y_h_pred, y_l, y_l_pred, y_c, y_c_pred, coeff_close = linear_regress(df)


            ## linear regression 1week

            month_1w_ago = dt.date.today() - dt.timedelta(7)
            month_1w_ago_ts = time.mktime(month_1w_ago.timetuple())
            df1w = df[(df.timestamp>month_1w_ago_ts)]
            xt1w, y_h1w, y_h_pred1w, y_l1w, y_l_pred1w, y_c1w, y_c_pred1w, coeff_close1w = linear_regress(df1w)


            ## linear regression 1month
            month_1_ago = monthdelta(dt.date.today(),-1)
            month_1_ago_ts = time.mktime(month_1_ago.timetuple())
            df1 = df[(df.timestamp>month_1_ago_ts)]
            xt1, y_h1, y_h_pred1, y_l1, y_l_pred1, y_c1, y_c_pred1, coeff_close1m = linear_regress(df1)


            ## linear regression 3months
            month_3_ago = monthdelta(dt.date.today(),-3)   #####------->  cambiato da 3 a 6 mesi!!!
            month_3_ago_ts = time.mktime(month_3_ago.timetuple())
            df3 = df[(df.timestamp>month_3_ago_ts)]
            xt3, y_h3, y_h_pred3, y_l3, y_l_pred3, y_c3, y_c_pred3, coeff_close3m = linear_regress(df3)

            ## linear regression 6months
            month_6_ago = monthdelta(dt.date.today(),-6)
            month_6_ago_ts = time.mktime(month_6_ago.timetuple())
            df6 = df[(df.timestamp>month_6_ago_ts)]
            xt6, y_h6, y_h_pred6, y_l6, y_l_pred6, y_c6, y_c_pred6, coeff_close6m = linear_regress(df6)


            #####################################  all months

            df['low_err'] = y_l-y_l_pred
            y_l_q = df.low_err.quantile(0.95)
            y_pred_l = y_l_pred-y_l_q           #--->
            df['y_pred_l'] = y_pred_l

            df['high_err'] = y_h-y_h_pred
            y_h_q = df.high_err.quantile(0.95)
            y_pred_h = y_h_pred+y_h_q           #--->
            df['y_pred_h'] = y_pred_h

            #####################################  3 months

            df3['low_err3'] = y_l3-y_l_pred3
            y_l_q3 = df3.low_err3.quantile(0.95)
            y_pred_l3 = y_l_pred3-y_l_q3           #--->
            df3['y_pred_l3'] = y_pred_l3

            df3['high_err3'] = y_h3-y_h_pred3
            y_h_q3 = df3.high_err3.quantile(0.95)
            y_pred_h3 = y_h_pred3+y_h_q3           #--->
            df3['y_pred_h3'] = y_pred_h3

            ##################################### esporto df all

            self.df_all = df.copy()

            ####   dati ultimo giorno    ###########################################################

            last_date = df.index[-1]
            last_ts = df.timestamp[-1]
            last_close = df.Close[-1]
            last_pred_l = df.y_pred_l[-1]
            last_pred_h = df.y_pred_h[-1]
            last_pred_c = y_c_pred[-1]

            last_pred_l3 = df3.y_pred_l3[-1]  #**3m
            last_pred_h3 = df3.y_pred_h3[-1]  #**3m

            #######################################

            ## creo data frame su una linea

            df_last = df.iloc[-1:,:].copy()#@

            ## creo colonna buy
            lower_gap = (last_close/last_pred_l)
            if lower_gap <1.1: df_last['buy'] = 'YES'
            if lower_gap >=1.1: df_last['buy'] = 'NO'

            # creo colonna buy3

            lower_gap3 = (last_close/last_pred_l3)
            if lower_gap3 <1.1: df_last['buy3'] = 'YES'
            if lower_gap3 >=1.1: df_last['buy3'] = 'NO'

            # creo le altre colonne:    attenzione 'gain' calcolato ipotizzando investimento 1000$

            df_last['lower_gap']=np.around(lower_gap,decimals=2)
            df_last['gain'] = int(((last_pred_h/last_close)-1)*1000)
            df_last['var_10gg']=round(((df.Close[-1:]/df.Close[-10:-1].min()).iat[0]*100)-100,2)
            df_last['var_1gg']=round(((df.Open[-1:]/df.Close[-2:-1].min()).iat[0]*100)-100,2)
            df_last['gain_ave'] = int(((last_pred_c/last_close)-1)*1000)
            df_last['stock'] = stock

            if coeff_close<0: df_last['trend'] = 0  #'FALLING'
            if coeff_close>0: df_last['trend'] = 'RAISING'

            df_last['slope_x1000'] = np.around(coeff_close*1000,decimals=5)
            df_last['first_data'] = df.index[0]
            df_last['name_ist'] = name_ist_str

            # creo le altre colonne ## 1w

            df_last['trend1w'] = coeff_close1w
            if coeff_close1w<0: df_last['trend1w'] = 0  #'FALLING'
            if coeff_close1w>0: df_last['trend1w'] = 1  #'RAISING'

            # creo le altre colonne ## 1m

            df_last['trend1m'] = coeff_close1m
            if coeff_close1m<0: df_last['trend1m'] = 0  #'FALLING'
            if coeff_close1m>0: df_last['trend1m'] = 1  #'RAISING'

            # creo le altre colonne ## 3m

            df_last['trend3m'] = coeff_close3m
            if coeff_close3m<0: df_last['trend3m'] = 0  #'FALLING'
            if coeff_close3m>0: df_last['trend3m'] = 1  #'RAISING'

            # creo le altre colonne ## 6m

            df_last['trend6m'] = coeff_close6m
            if coeff_close6m<0: df_last['trend6m'] = 0  #'FALLING'
            if coeff_close6m>0: df_last['trend6m'] = 1  #'RAISING'



            ######################################  esporto il df df_last

            self.df_last = df_last.copy()


            ###  per le def devo creare le variabili di istanza  ## esporto altri dati recuperbili dall'istanza

            self.xt = xt.copy()              #----> Datetime index
            self.y_c = y_c            #-----> Close column
            self.y_c_pred = y_c_pred  #-----> y_closure_prediction
            self.y_h = y_h            #------> High column
            self.y_pred_h = y_pred_h  #------> High prediction
            self.y_l = y_l            #------> Low column
            self.y_pred_l = y_pred_l  #------> Low prediction

            self.xt3 = xt3              #----> Datetime index
            self.y_c3 = y_c3            #-----> Close column
            self.y_c_pred3 = y_c_pred3  #-----> y_closure_prediction
            self.y_h3 = y_h3            #------> High column
            self.y_pred_h3 = y_pred_h3  #------> High prediction
            self.y_l3 = y_l3            #------> Low column
            self.y_pred_l3 = y_pred_l3  #------> Low prediction

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(self.stock,"Failed at 1°step because of ", e, exc_type, fname, exc_tb.tb_lineno)



    def ret_dataframe(self):
        try:
            return self.df_last
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(self.stock,"Failed at 1°step because of ", e, exc_type, fname, exc_tb.tb_lineno)





#################################################################################################################
#################################################################################################################
#################################################################################################################
#################################################################################################################
#######################################               START           ###########################################
#################################################################################################################
#################################################################################################################
#################################################################################################################
#################################################################################################################


start_time = time.time()

# no perché se la volta precedente alcuni stock non funzionano la volta successiva non li calcola
# infatti symb.txt viene aggiornata dopo la routine

try:
 with open('../symb.txt', "rb") as rb:
     symb = pickle.load(rb)
except:
 from list_stock import symb

from list_stock import symb



if len(symb)<1: from list_stock import symb


from list_stock_rem import symb_rem
from list_stock_new import symb_new

for i in symb_rem:
    try:
        symb.remove(i)
    except:
        pass

symb.extend(symb_new)


# # importa la lista ridotta x Test
# from list_stock import symb4
# symb = symb4



def Remove(duplicate):
    ''' rimuove i duplicati nella lista per evitare doppio calcolo'''
    final_list = []
    for num in duplicate:
        if num not in final_list:
            final_list.append(num)
    return final_list


symb=Remove(symb)


#
#

#########    CREA!!!!!          ###################
#**************************************************

df1 = pd.DataFrame({})
ist_dict = {}
i = 1
removed_list=[]
for stock in symb:
    name_ist = 'ist'+str(i)
    ist_dict[name_ist] = Yahoo_Scan(stock,'01/01/2006','d',name_ist)
    df1 = df1.append(ist_dict[name_ist].ret_dataframe(), sort=False)
    i += 1

#######  il dizionario ist_dict è quello con tutte le istanze
#######  il df "df1" è quello con gli indici: una riga per stock


##################################################################################



## CREO UN DIZIONARIO CON I NOMI DELLE ISTANZE ANZICHE IL NUMERO

stock_dict = {}

for i in ist_dict.keys():
    #print(i)
    #print(ist_dict[i].stock)

    stock_dict[ist_dict[i].stock] = ist_dict[i]

    # IL FILE CREATO E TROPPO GRANDE PER PYTHONANYWHERE
    #ist_dict[i].df_all.to_hdf('df_alltogheter.h5', key=ist_dict[i].stock)














##################################################################################


#######    database di output       ##########
df = df1.copy()

df_best = df[((df.buy3=='YES')&(
    df.buy=='YES')&(
    df.gain>100)&(
    df.slope_x1000>0.00001))]
df_best = df_best.sort_values(by=['gain'])

df_best_sup = df[((
    df.trend=='RAISING')&(
    df.var_1gg<-4)&(
    df.gain>100)&(
    df.slope_x1000>0.00001))]
df_best_sup = df_best_sup.sort_values(by=['gain'])


df_best_sup_10gg = df[((
    df.trend=='RAISING')&(
    df.var_10gg<-4)&(
    df.gain>100)&(
    df.slope_x1000>0.00001))]
df_best_sup = df_best_sup.sort_values(by=['gain'])


#  creo lista con tutti gli stock che sono passati e aggiorno il file

symb_pass_list = df1.stock.tolist()

try:
    with open('../symb.txt', "wb") as rb:
        pickle.dump(symb_pass_list, rb)
except:
    pass

# creo le diverse liste di stock in funzione delle performances


symb_pass_good = df_best.stock.tolist()
symb_pass_good_sup = df_best_sup.stock.tolist()
symb_pass_good_sup_10gg = df_best_sup_10gg.stock.tolist()


symb= df_best.stock.tolist()




####  SALVO I RISULTATI IN UN FILE #######################

# necessario from pandas import ExcelWriter
# richiede anche pip3 install openpyxl


nomedir = './{}'.format('DATA')
if os.path.exists(nomedir): shutil.rmtree(nomedir)
if not os.path.exists(nomedir): os.makedirs(nomedir)

nomefile= nomedir +'/'+ (str(dt.date.today())) + "best.csv"
export_csv = df.to_csv (nomefile, header=True, sep=';')

try:
    nomefile= nomedir +'/'+ (str(dt.date.today())) + "_best.xlsx"
    writer = pd.ExcelWriter(nomefile)


    df.sort_values('var_1gg').to_excel(writer,'data')
    writer.save()

    df_best.sort_values('var_1gg').to_excel(writer,'best')
    writer.save()

    df_best_sup.sort_values('var_1gg').to_excel(writer,'best_sup')
    writer.save()

    df_best_sup_10gg.sort_values('var_1gg').to_excel(writer,'best_sup_10gg')
    writer.save()
except:
    print("close the file xls")



def plot_ist(ist_in, working_folder):

    ist_corr = ist_dict[ist_in]

    plt.rcParams['figure.figsize'] = 18,5

    f = plt.figure()    # se uso questo non ho piu bisogno di inportare Figure
    a = f.add_subplot(111)

    a = plt.subplot2grid((9, 4), (0, 0), rowspan=5, colspan=4)
    a2 = plt.subplot2grid((9, 4), (5, 0), rowspan=2, colspan=4, sharex=a)
    a3 = plt.subplot2grid((9, 4), (7, 0), rowspan=2, colspan=4, sharex=a)


    a.plot(ist_corr.df_all.index,ist_corr.df_all.Close, label="price", linestyle='-', linewidth=1, color='b')
    a.plot(ist_corr.xt,ist_corr.y_c_pred, linestyle='--', color='r')
    a.plot(ist_corr.xt,ist_corr.y_pred_h, linestyle='--', color='gray')
    a.plot(ist_corr.xt,ist_corr.y_pred_l, linestyle='--', color='gray')

    a2.fill_between(ist_corr.df_all.index,ist_corr.df_all.Volume, 0)

    a3.plot(ist_corr.xt,ist_corr.df_all.macd)
    a3.plot(ist_corr.xt,ist_corr.df_all.signal)
    a3.bar(ist_corr.df_all.index,ist_corr.df_all.histogram,width = 3)

    a.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))

    plt.setp(a.get_xticklabels(), visible=False)
    plt.setp(a2.get_xticklabels(), visible=False)

    a.set_title(ist_corr.stock+' Max gain:'+
                str(ist_corr.df_last.gain.tolist()[0])+
                str('\$')+"  average gain: "+str(ist_corr.df_last.gain_ave.tolist()[0])+
                str('\$'))

    nomefile= working_folder + "/" + (str(dt.date.today())) + ist_corr.stock +".png"
    plt.savefig(nomefile)

    a.clear


# creo lista delle istanze best, plotto e salvo i file in PIC

list1=df_best.sort_values('stock').name_ist.astype(str).tolist()
working_folder = './{}'.format('PIC')
if os.path.exists(working_folder): shutil.rmtree(working_folder)
if not os.path.exists(working_folder): os.makedirs(working_folder)
for i in list1:
    plot_ist(i, working_folder)

# creo lista con tutte le istanze, plotto e salvo i file in PIC

list2=df.sort_values('var_1gg').name_ist.astype(str).tolist()[:15]
working_folder = './{}'.format('PIC2')
if os.path.exists(working_folder): shutil.rmtree(working_folder)
if not os.path.exists(working_folder): os.makedirs(working_folder)
for i in list2:
    plot_ist(i, working_folder)

# plotto il tempo che ci ha messo per girare
print("--- %s minutes ---" % ((time.time() - start_time)/60))


###############################################################################################################
################################    DOC WORD  AND HTML ########################################################
###############################################################################################################

document = Document ()

document.add_heading ( 'Buongiorno Fabrizio!\n the best Trade updated '+ str(dt.datetime.now()).split(".")[0] +'\n are ready!!!', level = 1 )


document.add_heading ( 'All The best trade' , level = 2 )

df_dall = df[['trend3m','trend','gain','stock','slope_x1000','name_ist','var_10gg','var_1gg']].sort_values('var_1gg')

t = document.add_table(df_dall.shape[0]+1, df_dall.shape[1])

# add the header rows.
for j in range(df_dall.shape[-1]):
    t.cell(0,j).text = df_dall.columns[j]

# add the rest of the data frame
for i in range(df_dall.shape[0]):
    for j in range(df_dall.shape[-1]):
        t.cell(i+1,j).text = str(df_dall.values[i,j])



document.add_heading ( 'The best trade' , level = 2 )

df_db = df_best[['trend3m','trend','gain','stock','slope_x1000','name_ist','var_10gg','var_1gg']].sort_values('stock')

t = document.add_table(df_db.shape[0]+1, df_db.shape[1])

# add the header rows.
for j in range(df_db.shape[-1]):
    t.cell(0,j).text = df_db.columns[j]

# add the rest of the data frame
for i in range(df_db.shape[0]):
    for j in range(df_db.shape[-1]):
        t.cell(i+1,j).text = str(df_db.values[i,j])



document.add_heading ( 'The best trade at 1gg' , level = 2 )

df_dbs = df_best_sup[['trend3m','trend','gain','stock','slope_x1000','name_ist','var_10gg','var_1gg']].sort_values('stock')

t = document.add_table(df_dbs.shape[0]+1, df_dbs.shape[1])

# add the header rows.
for j in range(df_dbs.shape[-1]):
    t.cell(0,j).text = df_dbs.columns[j]

# add the rest of the data frame
for i in range(df_dbs.shape[0]):
    for j in range(df_dbs.shape[-1]):
        t.cell(i+1,j).text = str(df_dbs.values[i,j])




document.add_heading ( 'The best trade at 10gg' , level = 2 )

df_dbs10 = df_best_sup_10gg[['trend3m','trend','gain','stock','slope_x1000','name_ist','var_10gg','var_1gg']].sort_values('stock')

t = document.add_table(df_dbs10.shape[0]+1, df_dbs10.shape[1])

# add the header rows.
for j in range(df_dbs10.shape[-1]):
    t.cell(0,j).text = df_dbs10.columns[j]

# add the rest of the data frame
for i in range(df_dbs10.shape[0]):
    for j in range(df_dbs10.shape[-1]):
        t.cell(i+1,j).text = str(df_dbs10.values[i,j])



###############  PICTURES ###############################

document . add_page_break ()

paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

document.add_heading ( 'Today highest drop start' , level = 2 )


#
files = []
filespath = []
# r=root, d=directories, f = files
for r, d, f in os.walk('../PIC2'):
    for file in f:
        if '.png' in file:
            files.append(file)
            filespath.append(os.path.join(r, file))

for file in files:
    document . add_picture ( './PIC2/'+file , width = Inches (7))

last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


















document . add_page_break ()

paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

document.add_heading ( 'The best trade graphs' , level = 2 )


#
files = []
filespath = []
# r=root, d=directories, f = files
for r, d, f in os.walk('../PIC'):
    for file in f:
        if '.png' in file:
            files.append(file)
            filespath.append(os.path.join(r, file))


for file in files:
    document . add_picture ( './PIC/'+file , width = Inches (7))



last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT




document.save ( 'todaytrade.docx' )



f = open("../todaytrade.docx", 'rb')
prev_folder = os.path.dirname(os.getcwd())
temp_folder = os.path.join(prev_folder,"mysite","templates", "../todaytrade.html")
print('Save html file in: ', temp_folder)
b = open(temp_folder, 'wb')
document = mammoth.convert_to_html(f)
b.write(document.value.encode('utf8'))
f.close()
b.close()