import os
import datetime as dt
import shutil
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mammoth
import sys
import feather
import h5py
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as mticker
import pickle
import warnings
warnings.filterwarnings('ignore')


def yeardelta(date,delta):
    '''
    :param date: data iniziale
    :param delta: data da aggiungere
    :return: fornisce la data + la data da aggiungere
    '''

    try:
        format_str = '%Y-%m-%d'
        date = dt.datetime.strptime(date, format_str)
    except:
        pass

    y1 = date.year+delta

    return date.replace(day=date.day, month=date.month, year=y1)

def monthdelta(date, delta):
    '''
    :param date: data iniziale
    :param delta: mesi da aggiungere
    :return: data iniziale + mesi da aggiungere = formato data
    '''

    try:
        format_str = '%Y-%m-%d'
        date = dt.datetime.strptime(date, format_str)
    except:
        pass

    m, y = (date.month+delta) % 12, date.year + ((date.month)+delta-1) // 12
    if not m: m = 12
    d = min(date.day, [31,
        29 if y%4==0 and not y%400==0 else 28,31,30,31,30,31,31,30,31,30,31][m-1])
    return date.replace(day=d,month=m, year=y)

def export_excel(df, df_best, df_best_sup, df_best_sup_10gg):
    # necessario from pandas import ExcelWriter
    # richiede anche pip3 install openpyxl


    nomedir = './{}'.format('DATA')
    if os.path.exists(nomedir): shutil.rmtree(nomedir)
    if not os.path.exists(nomedir): os.makedirs(nomedir)

    nomefile= nomedir +'/'+ (str(dt.date.today())) + "best.csv"
    export_csv = df.to_csv (nomefile, header=True, sep=';')

    try:
        nomefile= nomedir +'/'+ (str(dt.date.today())) + "_best.xlsx"
        writer = pd.ExcelWriter(nomefile)


        df.sort_values('var_1gg').to_excel(writer,'data')
        writer.save()

        df_best.sort_values('var_1gg').to_excel(writer,'best')
        writer.save()

        df_best_sup.sort_values('var_1gg').to_excel(writer,'best_sup')
        writer.save()

        df_best_sup_10gg.sort_values('var_1gg').to_excel(writer,'best_sup_10gg')
        writer.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print("Failed to create excel ", e, exc_type, fname, exc_tb.tb_lineno)
        print("Error maybe the xlsx file it is opened: close the file xls")

def export_hdf_feather(ist_dict):
    '''
    :param ist_dict: is a dictionary with the istance of the class
    :return: nothing just store files
    '''
    for i in ist_dict.keys():
        try:
            ist_dict[i].df_all.to_hdf('./DB/df.h5', key=ist_dict[i].stock, mode='w')
            # feather does not support serializing like datetime must use reset_index():
            ist_dict[i].df_all.reset_index().to_feather('./DB/'+ist_dict[i].stock+'.h5')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(ist_dict[i].stock, "Failed to store: ", e, exc_type, fname, exc_tb.tb_lineno)



def create_word_html(df, df_best, df_best_sup, df_best_sup_10gg):
    '''
    :param df:
    :param df_best:
    :param df_best_sup:
    :param df_best_sup_10gg:
    :return: nothing just save first a word and the convert it to html
    tobe improved --> use plotly and dynamic table for your web page
    '''

    document = Document()

    document.add_heading(
        'Buongiorno Fabrizio!\n the best Trade updated ' + str(dt.datetime.now()).split(".")[0] + '\n are ready!!!',
        level=1)

    document.add_heading('All The best trade', level=2)

    tablecolumns = ['datestamp','trend3m', 'trend', 'gain', 'stock', 'slope_x1000', 'name_ist', 'var_10gg', 'var_1gg',
                  'industry', 'volumechange', 'sentimentchange', 'wk52_high', 'mkt_Cap_bill','deltacovid','covidrecover']

    df_dall = df[tablecolumns].sort_values('var_1gg')
    ########         caution !!!!!!!!!!!!  ###################
    df_dall = df_dall = df[tablecolumns].sort_values('deltacovid')


    t = document.add_table(df_dall.shape[0] + 1, df_dall.shape[1])

    # add the header rows.
    for j in range(df_dall.shape[-1]):
        t.cell(0, j).text = df_dall.columns[j]

    # add the rest of the data frame
    for i in range(df_dall.shape[0]):
        for j in range(df_dall.shape[-1]):
            t.cell(i + 1, j).text = str(df_dall.values[i, j])

    document.add_heading('The best trade', level=2)

    df_db = df_best[tablecolumns].sort_values('stock')

    t = document.add_table(df_db.shape[0] + 1, df_db.shape[1])

    # add the header rows.
    for j in range(df_db.shape[-1]):
        t.cell(0, j).text = df_db.columns[j]

    # add the rest of the data frame
    for i in range(df_db.shape[0]):
        for j in range(df_db.shape[-1]):
            t.cell(i + 1, j).text = str(df_db.values[i, j])

    document.add_heading('The best trade at 1gg', level=2)

    df_dbs = df_best_sup[tablecolumns].sort_values('stock')

    t = document.add_table(df_dbs.shape[0] + 1, df_dbs.shape[1])

    # add the header rows.
    for j in range(df_dbs.shape[-1]):
        t.cell(0, j).text = df_dbs.columns[j]

    # add the rest of the data frame
    for i in range(df_dbs.shape[0]):
        for j in range(df_dbs.shape[-1]):
            t.cell(i + 1, j).text = str(df_dbs.values[i, j])

    document.add_heading('The best trade at 10gg', level=2)

    df_dbs10 = df_best_sup_10gg[tablecolumns].sort_values('stock')

    t = document.add_table(df_dbs10.shape[0] + 1, df_dbs10.shape[1])

    # add the header rows.
    for j in range(df_dbs10.shape[-1]):
        t.cell(0, j).text = df_dbs10.columns[j]

    # add the rest of the data frame
    for i in range(df_dbs10.shape[0]):
        for j in range(df_dbs10.shape[-1]):
            t.cell(i + 1, j).text = str(df_dbs10.values[i, j])



    ###############  PICTURES ###############################

    document . add_page_break ()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_heading ( 'Today highest drop start' , level = 2 )

    files = []
    filespath = []
    # r=root, d=directories, f = files
    for r, d, f in os.walk('./PIC2'):
        for file in f:
            if '.png' in file:
                files.append(file)
                filespath.append(os.path.join(r, file))

    for file in files:
        document . add_picture ( './PIC2/'+file , width = Inches (7))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document . add_page_break ()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_heading ( 'The best trade graphs' , level = 2 )

    files = []
    filespath = []
    # r=root, d=directories, f = files
    for r, d, f in os.walk('./PIC'):
        for file in f:
            if '.png' in file:
                files.append(file)
                filespath.append(os.path.join(r, file))

    for file in files:
        document . add_picture ( './PIC/'+file , width = Inches (7))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save ( 'todaytrade.docx' )

    f = open("todaytrade.docx", 'rb')
    prev_folder = os.path.dirname(os.getcwd())
    temp_folder = os.path.join(prev_folder,"mysite","templates","todaytrade.html")
    print('Html file daved in the folder : ', temp_folder)
    b = open(temp_folder, 'wb')
    document = mammoth.convert_to_html(f)
    b.write(document.value.encode('utf8'))
    f.close()
    b.close()

def plot_ist(ist_in, working_folder, ist_dict):
    '''
    Def create per generare le immagini che vanno inserite nel word e html
    :param ist_in: l'instanza di cui plottare i grafici
    :param working_folder: la cartella in cui salvare i graphici in .png
    :param ist_dict: è il dict con tutte le istanze
    :return: nothing just save picture files
    :import: matplotlib as plt
    '''


    ist_corr = ist_dict[ist_in]

    plt.rcParams['figure.figsize'] = 18,5

    f = plt.figure()    # se uso questo non ho piu bisogno di inportare Figure
    a = f.add_subplot(111)

    a = plt.subplot2grid((9, 4), (0, 0), rowspan=5, colspan=4)
    a2 = plt.subplot2grid((9, 4), (5, 0), rowspan=2, colspan=4, sharex=a)
    a3 = plt.subplot2grid((9, 4), (7, 0), rowspan=2, colspan=4, sharex=a)


    a.plot(ist_corr.df_all.index,ist_corr.df_all.Close, label="price", linestyle='-', linewidth=1, color='b')
    a.plot(ist_corr.xt,ist_corr.y_c_pred, linestyle='--', color='r')
    a.plot(ist_corr.xt,ist_corr.y_pred_h, linestyle='--', color='gray')
    a.plot(ist_corr.xt,ist_corr.y_pred_l, linestyle='--', color='gray')

    a2.fill_between(ist_corr.df_all.index,ist_corr.df_all.Volume, 0)

    a3.plot(ist_corr.xt,ist_corr.df_all.macd)
    a3.plot(ist_corr.xt,ist_corr.df_all.signal)
    a3.bar(ist_corr.df_all.index,ist_corr.df_all.histogram,width = 3)

    a.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))

    plt.setp(a.get_xticklabels(), visible=False)
    plt.setp(a2.get_xticklabels(), visible=False)

    a.set_title(ist_corr.stock+' Max gain:'+
                str(ist_corr.df_last.gain.tolist()[0])+
                str('\$')+"  average gain: "+str(ist_corr.df_last.gain_ave.tolist()[0])+
                str('\$'))

    nomefile= working_folder + "/" + (str(dt.date.today())) + ist_corr.stock +".png"
    plt.savefig(nomefile)

    a.clear

def make_folder(working_folder):
    '''
    set working folder if exist remove and create if not exists create
    :param working_folder: the folder
    :return: the working_folder itself
    '''
    if os.path.exists(working_folder): shutil.rmtree(working_folder)
    if not os.path.exists(working_folder): os.makedirs(working_folder)
    return working_folder

def save_pass_list(df=0, symb_pass=0):
    '''
    Se inserisco df1 prende la lista dei simboli del dataframe
    Posso invece direttamente inserire symb_pass
    :param df1:
    :param symb_pass:
    :return:
    '''
    if symb_pass==0:
        symb_pass = df.stock.tolist()
        try:
            with open('../symb.txt', "wb") as rb:
                pickle.dump(symb_pass, rb)
        except:
            print('symb pass list not saved')
    if symb_pass != 0:
        try:
            with open('../symb.txt', "wb") as rb:
                pickle.dump(symb_pass, rb)
        except:
            print('symb pass list not saved')

def retrieve_pass_list():
    try:
        with open('symb.txt', "rb") as rb:
            symb = pickle.load(rb)
        return symb
    except:
        print('symb_pass list not retrieved')





def retrieve_symb_list():
    '''
    questa funzione recupera gli stock da calcolare
    :return: lista di stock
    '''

    from list_stock import symb

    if len(symb) < 1: from list_stock import symb

    from list_stock_rem import symb_rem
    from list_stock_new import symb_new

    for i in symb_rem:
        try:
            symb.remove(i)
        except:
            pass

    symb.extend(symb_new)

    def Remove(duplicate):
        final_list = []
        for num in duplicate:
            if (num not in final_list) & (num not in symb_rem):
                final_list.append(num)
        return final_list

    symb = Remove(symb)
    symb.sort()

    print('The final symb list has length',len(symb),'\n and is :\n', symb)

    return symb